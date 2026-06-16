from __future__ import annotations

import html
import logging
from pathlib import Path

import fitz  # PyMuPDF
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document

from app.core.schemas import PageText

logger = logging.getLogger(__name__)


class DocumentLoaderError(RuntimeError):
    pass


def load_document(path: Path) -> list[PageText]:
    """Load a supported document into page-like text units."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix in {".txt", ".md"}:
        return _load_text(path)
    if suffix in {".csv"}:
        return _load_csv(path)
    if suffix in {".xlsx", ".xls"}:
        return _load_excel(path)
    if suffix in {".html", ".htm"}:
        return _load_html(path)
    raise DocumentLoaderError(f"Unsupported file type: {path.suffix}")


def _load_pdf(path: Path) -> list[PageText]:
    pages: list[PageText] = []
    try:
        with fitz.open(path) as doc:
            for i, page in enumerate(doc, start=1):
                text = page.get_text("text") or ""
                text = _clean_text(text)
                if text:
                    pages.append(PageText(text=text, page=i, metadata={"loader": "pymupdf"}))
    except Exception as exc:  # pragma: no cover - external parser errors
        raise DocumentLoaderError(f"Failed to parse PDF {path}: {exc}") from exc
    return pages


def _load_docx(path: Path) -> list[PageText]:
    try:
        doc = Document(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return [PageText(text=_clean_text("\n".join(parts)), page=None, metadata={"loader": "python-docx"})]
    except Exception as exc:  # pragma: no cover
        raise DocumentLoaderError(f"Failed to parse DOCX {path}: {exc}") from exc


def _load_text(path: Path) -> list[PageText]:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            text = path.read_text(encoding=encoding)
            return [PageText(text=_clean_text(text), page=None, metadata={"loader": "text", "encoding": encoding})]
        except UnicodeDecodeError:
            continue
    raise DocumentLoaderError(f"Failed to decode text file: {path}")


def _load_csv(path: Path) -> list[PageText]:
    try:
        df = pd.read_csv(path)
        text = df.to_markdown(index=False)
        return [PageText(text=_clean_text(text), page=None, metadata={"loader": "pandas-csv"})]
    except Exception as exc:  # pragma: no cover
        raise DocumentLoaderError(f"Failed to parse CSV {path}: {exc}") from exc


def _load_excel(path: Path) -> list[PageText]:
    try:
        sheets = pd.read_excel(path, sheet_name=None)
        parts: list[str] = []
        for name, df in sheets.items():
            parts.append(f"# Sheet: {name}\n" + df.to_markdown(index=False))
        return [PageText(text=_clean_text("\n\n".join(parts)), page=None, metadata={"loader": "pandas-excel"})]
    except Exception as exc:  # pragma: no cover
        raise DocumentLoaderError(f"Failed to parse Excel {path}: {exc}") from exc


def _load_html(path: Path) -> list[PageText]:
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.extract()
        text = html.unescape(soup.get_text("\n"))
        return [PageText(text=_clean_text(text), page=None, metadata={"loader": "beautifulsoup"})]
    except Exception as exc:  # pragma: no cover
        raise DocumentLoaderError(f"Failed to parse HTML {path}: {exc}") from exc


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r", "\n").split("\n")]
    compact: list[str] = []
    blank = False
    for line in lines:
        if not line:
            if not blank:
                compact.append("")
            blank = True
        else:
            compact.append(line)
            blank = False
    return "\n".join(compact).strip()
